# 1) Создайте класс для банковского счета с методами пополнения, снятия и проверки баланса.
class BankAccount:
    def __init__(self, name, balance):
        self.name = name
        self.balance = balance

    def deposit(self, amount):
        self.balance += amount
        print(f"{amount} has been deposited to {self.name}'s account")

    def withdraw(self, amount):
        if amount > self.balance:
            print(f"Insufficient balance in {self.name}'s account")
        else:
            self.balance -= amount
            print(f"{amount} has been withdrawn from {self.name}'s account")

    def check_balance(self):
        print(f"{self.name}'s current balance is {self.balance}")

account1 = BankAccount("Charlie Hudson", 5000)
account1.check_balance()
account1.deposit(2000)
account1.check_balance()
account1.withdraw(3000)
account1.check_balance()

# 2) Создайте класс для автомобиля с методами запуска, остановки и подачи сигнала.
class Car:
    def __init__(self, make, model, year):
        self.make = make
        self.model = model
        self.year = year
        self._speed = 0

    def accelerate(self):
        self._speed += 50

    def brake(self):
        self._speed -= 30

    def get_signal(self):
        print("Horn")

auto = Car("Audi", "Q3", 2020)
auto.get_signal()

# 3) Создайте класс для человека с атрибутами для имени, возраста и адреса.
class Person:
    def __init__(self, name, age, address):
        self.name = name
        self.age = age
        self.address = address

# 4) Создайте класс для университета с атрибутами имени, адреса и студентов.
class University:
    def __init__(self, name, address, students):
        self.name = name
        self.address = address
        self.students = students

# 5) Создайте класс для прямоугольника с атрибутами длины и ширины и методами вычисления площади и периметра.
class Rectangle:
    def __init__(self, length, width):
        self.length = length
        self.width = width

    def calculate_area(self):
        return self.length * self.width

    def calculate_perimeter(self):
        return 2 * (self.length + self.width)

rect1 = Rectangle(10, 15)
print(rect1.calculate_area())
print(rect1.calculate_perimeter())

# 6) Создайте класс для корзины покупок с методами добавления и удаления товаров и расчета общей стоимости.
from datetime import date,timedelta

class Product:

    def __init__(self, name, price, real_date, exp_date, brand="R"):
        self.name = name
        self.price = price
        self.brand = brand
        self.real_date = real_date
        self.exp_date = exp_date

class ProductList:

    def __init__(self, owner, date_now):
        self.L = list()
        self.owner = owner
        self.date_now = date_now

    def add(self, *args: Product):
        for arg in args:
            if arg.real_date + timedelta(days=arg.exp_date) >= self.date_now:
                self.L.append(arg)
            else:
                print("Expired")

    def __str__(self):
        p_str = ""
        for p in self.L:
            p_str = p.name + " "
        return f"{self.owner}. Your products: {p_str}"

    def remove_one(self, p_name):
        for i in range(len(self.L)):
            if self.L[i].name == p_name:
                self.L.pop(i)
                break

    def remove_all(self, p_name):
        for i in range(len(self.L)-1, -1, -1):
            if self.L[i].name == p_name:
                self.L.pop(i)

    def total(self):
        count = 0
        for p in self.L:
            count += p.price
        return count

ice_cream = Product("Ice cream", 90, date(2023, 4, 1), exp_date=120)
milk = Product("Milk", 120, date(2023, 4, 7), exp_date=5)
chocolate = Product("Chocolate", 85, date(2023, 3, 5), exp_date=100)

p1 = ProductList("Anna", date_now=date(2023, 4, 6))
p1.add(ice_cream, ice_cream, chocolate)
print(p1.total())
p1.remove_all("Ice cream")
print(p1)
print(p1.total())

# 7) Создайте класс для игры с атрибутами для имени игрока, очков, жизней и уровня. Класс должен содержать методы для:
# a)	Запуска, завершения и перезапуска игры
# b)	Получения очков (при наборе определённого количества очков уровень должен увеличиться на 1)
# c)	Потери жизней (при достижении нуля жизней должна появиться надпись Game Over). При перезапуске игры жизни
# восстанавливаются.
class Game:
    def __init__(self, name, score, level, lives=3):
        self.name = name
        self.score = score
        self.lives = lives
        self.level = level

    def run(self):
        return self.run

    def stop(self):
        return self.stop

    def rerun(self):
        return self.rerun

    def get_score(self, amount):
        self.score += amount
        if self.score == 10:
            self.level += 1

    def lose_lives(self):
        if self.lives == 0:
            print("Game over")

    def update_level(self):
        if self.rerun:
            self.lives = 3

# 8) Задание по библиотекам docx и python-docx (пример итогового файла в demo.docx): написать код, который будет
# создавать файл Word, где обязательно должны быть:
# a)	Заголовки двух или трёх разных уровней
# b)	Обычный текст, полужирный шрифт и курсив
# c)	Подчёркнутый текст
# d)	Нумерованный или ненумерованный список
# e)	Изображения
from docx import Document
from docx.shared import Inches

document = Document()
h = document.add_heading('Татьяна Владимировна Черниговская', 0)
p = document.add_paragraph('Российский биолог, лингвист, семиотик и психолог')
p2 = document.add_paragraph('Доктор биологических наук, доктор филологических наук, член-корреспондент РАО. \
Заслуженный деятель высшего образования и Заслуженный деятель науки РФ.')
run = p2.runs[0]
run.underline = True
h1 = document.add_heading("Лекции", 1)
run1 = h1.runs[0]
run1.bold = True
p3 = document.add_paragraph("Как научить мозг учиться")
p3.style = 'List Bullet'
run2 = p3.runs[0]
run2.italic = True
p4 = document.add_paragraph("Куда мы попали?")
p4.style = 'List Bullet'
run3 = p4.runs[0]
run3.italic = True
p5 = document.add_paragraph("Язык, мозг и гены")
p5.style = 'List Bullet'
run4 = p5.runs[0]
run4.italic = True
picture = document.add_picture("tvch.jpg", width=Inches(3))
picture1 = document.add_picture("tvch1.jpg", width=Inches(3))
picture2 = document.add_picture("tvch2.jpg", width=Inches(3))
document.save('My document.docx')

# 9) Задание по openpyxl: с помощью openpyxl написать код, который будет создавать какую-либо матрицу товаров (пример в
# stationery.xlsx) и в последнем столбце автоматически рассчитывать итоговую стоимость для каждой строки (кол-во
# товара*цену за единицу).
from openpyxl import Workbook

wb = Workbook()
ws = wb.active
wb.create_sheet("Sheet")
data = (
    ("Название", "Артикул", "Количество", "Цена", "Общая стоимость"),
    ("Карандаш простой", "a859", 5, 22.8),
    ("Ручка шариковая", "b125", 5, 27.5),
    ("Тетрадь 12 л.", "d385", 8, 15),
    ("Тетрадь 48 л.", "d399", 10, 65.2),
    ("Дневник школьный", "c764", 1, 230),
    ("Пенал", "f637", 1, 198.9),
    ("Карандаши цветные", "a421", 2, 87.6)
)
for row in data:
    ws.append(row)
for row in range(2, ws.max_row + 1):
    count = ws.cell(row, 3)
    price = ws.cell(row, 4)
    total_price = count.value * price.value
    total_price_cell = ws.cell(row, 5)
    total_price_cell.value = total_price
wb.save('items.xlsx')
wb.close()
